from docx import Document 


class DocxExtractor:
    def __init__(self):
        pass

    def extract(self, file_path):
        doc = Document(file_path)
        all_slices = []
        all_paragraph = 0
        all_table = 0
        index = 1

        for element in doc.element.body:
            if element.tag.endswith('sectPr'):
                continue
            if element.tag.endswith('p'):
                paragraph = doc.paragraphs[all_paragraph]
                paragraph_type = paragraph.style.name
                if "附录" in paragraph_type:
                    paragraph_type = "附录"
                if paragraph_type == "附录":
                    paragraph_content = paragraph_type + paragraph.text.strip()
                else:
                    paragraph_content = paragraph.text.strip()
                if len(paragraph_content.strip()) > 3:
                    all_slices.append(
                        {
                            'type':"paragraph",
                            "content":paragraph_content,
                            "paragraph_type":paragraph_type,
                            "pages":[index]
                        }
                    )
                else:
                    pass
                all_paragraph += 1
            elif element.tag.endswith('tbl'):
                table = doc.tables[all_table]
                table_data = ""
                for row in table.rows:
                    row_data = ""
                    for cell in row.cells:
                        row_data += "|" + str(cell.text.replace("\n", " "))
                    if len(table_data) != 0:
                        table_data += "\n"
                    table_data  += row_data[1:]
                str_table = str(table_data)
                for item in str_table:
                    if item not in ['[', ']', ',', ' ', "'"] and len(str_table)>2 and "|" in item:
                        all_slices.append(
                            {
                                'type':"table",
                                "content":str_table,
                                "paragraph_type":"table",
                                "pages":[index]
                            }
                        )
                        break
                    else:
                        continue
        return all_slices